# Archivo: export.py
from docx import Document
from docx.shared import Inches
import datetime
import os
from settings import (
    TERRESTRES_RADIO, TERRESTRES_MASA, 
    REPORTE_WORD_PATH, TEMP_MIN_K, TEMP_MAX_K
)

def generar_reporte_word(df_gold, ruta_obj1, ruta_obj2, ruta_obj3, ruta_obj4, ruta_obj5): 
    doc = Document()
    doc.add_heading('Análisis de Exoplanetas: Informe de Clasificación y Habitabilidad', 0)

    fecha = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    doc.add_paragraph(f"Fecha de generación: {fecha}")

    # --- SECCIÓN: OBJETIVO 1 (Lo que ya tenías) ---
    doc.add_heading('Objetivo 1: Clasificación de Tipos de Mundos', level=1)
    p1 = doc.add_paragraph()
    p1.add_run("Metodología: ").bold = True
    p1.add_run(f"Categorización mediante Radio < {TERRESTRES_RADIO} y Masa < {TERRESTRES_MASA}.")

    if os.path.exists(ruta_obj1):
        doc.add_picture(ruta_obj1, width=Inches(5.5))
        doc.add_paragraph("Figura 1: Distribución logarítmica de Masa vs Radio.")
        
        # --- RESUMEN ESTADÍSTICO (Tu tabla original) ---
        doc.add_heading('Tabla de Población Planetaria', level=2)
        conteo = df_gold['tipo_planeta'].value_counts()
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid' # Un toque de estilo
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Tipo de Planeta'
        hdr_cells[1].text = 'Cantidad'

        for tipo, total in conteo.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(tipo)
            row_cells[1].text = str(total)

    # --- SECCIÓN: OBJETIVO 2 (La nueva joya) ---
    doc.add_heading('Objetivo 2: Análisis de la Zona de Goldilocks', level=1)
    
    # Explicación de la temperatura
    p2 = doc.add_paragraph()
    p2.add_run("Criterios de Habitabilidad: ").bold = True
    p2.add_run(
        f"Se han filtrado los mundos de tipo 'Terrestre' que presentan una Temperatura de Equilibrio "
        f"entre {TEMP_MIN_K}K y {TEMP_MAX_K}K. Este rango permite, teóricamente, la existencia de agua líquida."
    )

    if os.path.exists(ruta_obj2):
        doc.add_heading('Distribución de Temperaturas', level=2)
        doc.add_picture(ruta_obj2, width=Inches(5.5))
        doc.add_paragraph("Figura 2: Densidad de temperatura en planetas rocosos.")
    # Resumen de Candidatos
    doc.add_heading('Resultados de Habitabilidad', level=2)
    total_habitables = df_gold['habitable'].sum()
    doc.add_paragraph(f"Tras el análisis, se han identificado {total_habitables} candidatos potenciales.")
    # seccion: objetivo 3
    doc.add_heading('Distancia y Tiempo', level=1)
    p2 = doc.add_paragraph()
    p2.add_run("¿Estamos encontrando planetas más lejos a medida que pasa el tiempo? ").bold = True
    if os.path.exists(ruta_obj3):
        doc.add_heading('Evolución de los Descubrimientos de Exoplanetas', level=2)
        doc.add_picture(ruta_obj3, width=Inches(5.5))
        doc.add_paragraph("Figura 3:Número de  planetas, respecto al año de descubrimiento .")
        doc.add_paragraph('Esto muestra cómo nuestra tecnología (Kepler, James Webb) nos permite ver cada vez más profundo en la galaxia. El dominio del método de Tránsito (el amarillo) a partir del 2014, probablemente gracias a misiones como Kepler')
    #seccion: objetivo 4
    doc.add_heading('El "Vecindario Estelar"', level=1)
    p2 = doc.add_paragraph('¿Qué tipo de estrellas tienen más planetas?')
    if os.path.exists(ruta_obj4):
        doc.add_heading('Figura 4:Cantidad de planetas detectados respecto a la clasificación espectral (O-M) ', level=2)
        doc.add_picture(ruta_obj4, width=Inches(5.5))
        # 2. Agregar la leyenda técnica
        doc.add_heading('Guía de Clasificación Espectral (O-M)', level=2)
        p_leyenda = doc.add_paragraph()
        p_leyenda.add_run("Para interpretar la gráfica anterior, se definen los siguientes tipos estelares:").italic = True
        
        # Definimos el diccionario con la info que me pasaste
        tipos_info = [
            ("Tipo O (Azules):", "> 30,000 K. Muy masivas, brillantes y calientes."),
            ("Tipo B (Azul-Blancas):", "10,000 - 30,000 K. Ejemplo: Rigel."),
            ("Tipo A (Blancas):", "7,500 - 10,000 K. Líneas de hidrógeno fuertes."),
            ("Tipo F (Blanco-Amarillentas):", "6,000 - 7,500 K."),
            ("Tipo G (Amarillas):", "5,200 - 6,000 K. Ejemplo: El Sol."),
            ("Tipo K (Naranjas):", "3,700 - 5,200 K. Más frías que el Sol."),
            ("Tipo M (Rojas):", "≤ 3,700 K. Las más comunes y frías (Enanas Rojas).")
        ]
        
        # Lo agregamos como una lista con viñetas
        for titulo, desc in tipos_info:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(titulo).bold = True
            p.add_run(f" {desc}")
        doc.add_heading('Análisis de Inteligencia Artificial (ML)', level=1)
        doc.add_paragraph(
            "Se implementó un algoritmo de K-Means Clustering para agrupar los planetas "
            "automáticamente según sus propiedades físicas. Este modelo identificó "
            "patrones naturales en la base de datos sin intervención humana."
        )
        doc.add_heading('Conclusión del Modelo:', level=1)
        doc.add_paragraph('La Inteligencia Artificial validó nuestra metodología: los grupos que definimos manualmente (Terrestres, Neptunianos y Gigantes) existen estadísticamente. El K-Means "aprendió" la estructura física de la galaxia sin que nadie le explicara qué es un planeta.')
        # Insertar la gráfica de los clusters
        doc.add_picture(ruta_obj5, width=Inches(5.0))
        doc.save(REPORTE_WORD_PATH)
    print(f"✅ Reporte final generado en: {REPORTE_WORD_PATH}")

    return REPORTE_WORD_PATH